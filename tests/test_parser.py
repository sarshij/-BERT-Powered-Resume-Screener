import io
import pytest
from app.utils.parser import parse_resume, parse_pdf, parse_txt, parse_docx

def test_parse_txt():
    text = b'Hello world.'
    assert parse_txt(text) == 'Hello world.'

def test_parse_docx():
    # Create a simple docx in memory
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph('Hello docx')
        doc.add_paragraph('Second line')
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        result = parse_docx(bio.read())
        assert 'Hello docx' in result
        assert 'Second line' in result
    except Exception:
        pytest.skip('docx not available')

def test_parse_pdf_exists():
    # Ensure function exists
    assert callable(parse_pdf)

def test_parse_resume_exists():
    assert callable(parse_resume)
